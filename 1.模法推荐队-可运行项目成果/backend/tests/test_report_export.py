"""Tests for report export to DOCX and PDF formats."""

from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services.report_exporter import get_report_exporter
from app.services.report_service import ReportGenerationService

client = TestClient(app)


def _make_report():
    svc = ReportGenerationService()
    return svc.generate(
        request_id="test-001",
        parse_result={
            "raw_text": "农户小额贷款贷前风控",
            "intent": "credit_risk",
            "business_scenario": "农户小额贷款贷前准入",
            "tags": ["admission_scoring"],
        },
        recommend_result={
            "recommendations": [
                {
                    "rank": 1,
                    "model_name": "准入评分卡",
                    "model_id": "RISK_001",
                    "total_score": 88.5,
                    "recommendation_reason": "场景匹配高",
                }
            ]
        },
    )


def test_exporter_to_docx():
    report = _make_report()
    content = get_report_exporter().to_docx(report)
    assert isinstance(content, bytes)
    assert len(content) > 1000
    # DOCX is a ZIP archive — magic bytes "PK"
    assert content[:2] == b"PK"

    rendered = Document(BytesIO(content))
    visible_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in rendered.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "**" not in visible_text
    assert "###" not in visible_text
    assert "88.5" not in visible_text + table_text
    assert len(rendered.tables) >= 2


def test_exporter_to_pdf():
    report = _make_report()
    content = get_report_exporter().to_pdf(report)
    assert isinstance(content, bytes)
    assert len(content) > 500
    # PDF magic bytes "%PDF"
    assert content[:4] == b"%PDF"


def test_export_endpoint_docx():
    resp = client.post(
        "/api/v1/reports/recommendation/export?format=docx",
        json={
            "request_id": "test-001",
            "demand_raw": "农户小额贷款贷前风控",
            "parse_result": {
                "raw_text": "农户小额贷款贷前风控",
                "intent": "credit_risk",
                "business_scenario": "农户小额贷款贷前准入",
                "tags": ["admission_scoring"],
            },
            "recommend_result": {
                "recommendations": [
                    {
                        "rank": 1,
                        "model_name": "准入评分卡",
                        "model_id": "RISK_001",
                        "total_score": 88.5,
                        "recommendation_reason": "场景匹配高",
                    }
                ]
            },
        },
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats"
    )
    content = resp.content
    assert len(content) > 1000
    assert content[:2] == b"PK"
    assert "attachment" in resp.headers.get("content-disposition", "")


def test_export_endpoint_pdf():
    resp = client.post(
        "/api/v1/reports/recommendation/export?format=pdf",
        json={
            "request_id": "test-002",
            "demand_raw": "县域新客首贷营销",
            "parse_result": {
                "raw_text": "县域新客首贷营销",
                "intent": "customer_marketing",
                "business_scenario": "县域新客首贷营销",
            },
            "recommend_result": {
                "recommendations": [
                    {
                        "rank": 1,
                        "model_name": "新客响应模型",
                        "model_id": "MKT_001",
                        "total_score": 85.0,
                        "recommendation_reason": "营销匹配",
                    }
                ]
            },
        },
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith("application/pdf")
    content = resp.content
    assert len(content) > 500
    assert content[:4] == b"%PDF"
    assert "attachment" in resp.headers.get("content-disposition", "")


def test_export_endpoint_invalid_format():
    resp = client.post(
        "/api/v1/reports/recommendation/export?format=xls",
        json={"demand_raw": "测试"},
    )
    assert resp.status_code == 400
