"""Render recommendation reports as structured DOCX and PDF documents."""

from __future__ import annotations

import html
import io
import re
from dataclasses import dataclass
from typing import Iterable

from app.schemas.report import ReportResponse


BOUNDARY_NOTICE = (
    "本报告基于当前会话输入、系统模型资产与推荐排序结果与可核验依据生成，仅供模型选型和人工复核；"
    "不构成自动化生产决策、真实经营收益承诺或银行模型市场已联调证明。"
)


@dataclass(frozen=True)
class _Block:
    kind: str
    value: str | list[list[str]]


def _clean_inline(text: str) -> str:
    text = text.replace("⚠", "缺失：").replace("✅", "已满足：")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.strip()


def _parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [_clean_inline(cell) for cell in line.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows


def _parse_blocks(content: str) -> Iterable[_Block]:
    lines = content.splitlines()
    paragraph: list[str] = []

    def flush_paragraph() -> Iterable[_Block]:
        if paragraph:
            value = " ".join(item.strip() for item in paragraph if item.strip())
            paragraph.clear()
            if value:
                yield _Block("paragraph", value)

    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            yield from flush_paragraph()
            index += 1
            continue
        if line.startswith("|") and line.endswith("|"):
            yield from flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                index += 1
            rows = _parse_table(table_lines)
            if rows:
                yield _Block("table", rows)
            continue
        heading = re.match(r"^#{2,4}\s+(.+)$", line)
        if heading:
            yield from flush_paragraph()
            yield _Block("heading", _clean_inline(heading.group(1)))
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            yield from flush_paragraph()
            yield _Block("bullet", _clean_inline(bullet.group(1)))
            index += 1
            continue
        numbered = re.match(r"^\d+[.、]\s*(.+)$", line)
        if numbered:
            yield from flush_paragraph()
            yield _Block("number", _clean_inline(numbered.group(1)))
            index += 1
            continue
        paragraph.append(line)
        index += 1
    yield from flush_paragraph()


def _inline_xml(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"`([^`]*)`", r"<font name='Courier'>\1</font>", escaped)
    return escaped


class ReportExporter:
    """Convert a ReportResponse into readable, printable binary formats."""

    def to_docx(self, report: ReportResponse) -> bytes:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor

        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.start_type = WD_SECTION.NEW_PAGE

        styles = doc.styles
        for style_name, size in (("Normal", 10.5), ("Title", 21), ("Heading 1", 15), ("Heading 2", 12)):
            style = styles[style_name]
            style.font.name = "Microsoft YaHei"
            style.font.size = Pt(size)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        styles["Normal"].paragraph_format.space_after = Pt(5)
        styles["Normal"].paragraph_format.line_spacing = 1.25
        styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
        styles["Heading 1"].paragraph_format.space_before = Pt(12)
        styles["Heading 1"].paragraph_format.space_after = Pt(5)
        styles["Heading 2"].font.color.rgb = RGBColor(47, 84, 150)

        title = doc.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title.add_run(report.title or "模型推荐报告")

        if report.summary:
            summary = doc.add_paragraph()
            summary.paragraph_format.space_after = Pt(8)
            run = summary.add_run(report.summary)
            run.bold = True
            run.font.color.rgb = RGBColor(55, 65, 81)

        notice = doc.add_paragraph()
        notice.paragraph_format.space_after = Pt(10)
        notice_run = notice.add_run("证据与使用边界：" + BOUNDARY_NOTICE)
        notice_run.font.size = Pt(9)
        notice_run.font.color.rgb = RGBColor(146, 64, 14)

        for report_section in report.sections:
            blocks = list(_parse_blocks(report_section.content))
            first_table = next((block for block in blocks if block.kind == "table"), None)
            if first_table is not None and isinstance(first_table.value, list):
                if first_table.value and max(len(row) for row in first_table.value) == 5:
                    doc.add_page_break()
                    doc.add_paragraph("")
            doc.add_heading(report_section.title, level=1)
            number_index = 0
            for block in blocks:
                if block.kind == "heading":
                    doc.add_heading(str(block.value), level=2)
                elif block.kind == "bullet":
                    doc.add_paragraph(str(block.value), style="List Bullet")
                elif block.kind == "number":
                    number_index += 1
                    doc.add_paragraph(f"{number_index}. {block.value}")
                elif block.kind == "table":
                    rows = block.value
                    assert isinstance(rows, list)
                    column_count = max(len(row) for row in rows)
                    table = doc.add_table(rows=len(rows), cols=column_count)
                    table.style = "Table Grid"
                    table.autofit = column_count != 5
                    column_widths = [Cm(1), Cm(3.2), Cm(2.2), Cm(1.4), Cm(9)] if column_count == 5 else None
                    if column_widths:
                        for col_index, width in enumerate(column_widths):
                            table.columns[col_index].width = width
                    for row_index, row in enumerate(rows):
                        for col_index, value in enumerate(row):
                            cell = table.cell(row_index, col_index)
                            if column_widths:
                                cell.width = column_widths[col_index]
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.text = value
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.space_after = Pt(0)
                                for run in paragraph.runs:
                                    run.font.name = "Microsoft YaHei"
                                    run.font.size = Pt(8 if len(rows[0]) >= 5 else 9)
                                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                                    run.bold = row_index == 0
                    for table_row in table.rows:
                        cannot_split = OxmlElement("w:cantSplit")
                        table_row._tr.get_or_add_trPr().append(cannot_split)
                    header_row = table.rows[0]._tr
                    table_header = OxmlElement("w:tblHeader")
                    table_header.set(qn("w:val"), "true")
                    header_row.get_or_add_trPr().append(table_header)
                    doc.add_paragraph("")
                else:
                    self._add_docx_inline_paragraph(doc, str(block.value))

        doc.add_heading("报告元数据", level=1)
        metadata = doc.add_table(rows=3, cols=2)
        metadata.style = "Table Grid"
        for index, (label, value) in enumerate((
            ("报告编号", report.report_id),
            ("生成时间", report.generated_at),
            ("生成来源", report.generation_source),
        )):
            metadata.cell(index, 0).text = label
            metadata.cell(index, 1).text = value

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("银行模型市场智能推荐助手 | 仅供人工复核")
        footer_run.font.name = "Microsoft YaHei"
        footer_run.font.size = Pt(8)
        footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _add_docx_inline_paragraph(doc: object, text: str) -> None:
        paragraph = doc.add_paragraph()
        cursor = 0
        for match in re.finditer(r"\*\*(.*?)\*\*", text):
            if match.start() > cursor:
                paragraph.add_run(text[cursor : match.start()])
            run = paragraph.add_run(match.group(1))
            run.bold = True
            cursor = match.end()
        if cursor < len(text):
            paragraph.add_run(text[cursor:])

    def to_pdf(self, report: ReportResponse) -> bytes:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import (
            KeepTogether,
            ListFlowable,
            ListItem,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        try:
            pdfmetrics.getFont("STSong-Light")
        except KeyError:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        font_name = "STSong-Light"
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ChineseTitle", parent=styles["Title"], fontName=font_name,
            fontSize=20, leading=27, textColor=colors.HexColor("#1F4E79"),
            spaceAfter=10 * mm, wordWrap="CJK",
        )
        summary_style = ParagraphStyle(
            "ChineseSummary", parent=styles["BodyText"], fontName=font_name,
            fontSize=10.5, leading=17, textColor=colors.HexColor("#374151"),
            backColor=colors.HexColor("#F3F6FA"), borderPadding=7,
            spaceAfter=5 * mm, wordWrap="CJK",
        )
        notice_style = ParagraphStyle(
            "ChineseNotice", parent=styles["BodyText"], fontName=font_name,
            fontSize=8.5, leading=14, textColor=colors.HexColor("#92400E"),
            backColor=colors.HexColor("#FFF7ED"), borderPadding=6,
            spaceAfter=5 * mm, wordWrap="CJK",
        )
        h1_style = ParagraphStyle(
            "ChineseH1", parent=styles["Heading1"], fontName=font_name,
            fontSize=14, leading=19, textColor=colors.HexColor("#1F4E79"),
            spaceBefore=5 * mm, spaceAfter=2.5 * mm, wordWrap="CJK",
        )
        h2_style = ParagraphStyle(
            "ChineseH2", parent=styles["Heading2"], fontName=font_name,
            fontSize=11.5, leading=16, textColor=colors.HexColor("#2F5496"),
            spaceBefore=3 * mm, spaceAfter=1.5 * mm, wordWrap="CJK",
        )
        body_style = ParagraphStyle(
            "ChineseBody", parent=styles["BodyText"], fontName=font_name,
            fontSize=9.5, leading=15, spaceAfter=2.2 * mm, wordWrap="CJK",
        )
        small_style = ParagraphStyle(
            "ChineseSmall", parent=body_style, fontSize=7.5, leading=11,
            spaceAfter=0, wordWrap="CJK",
        )

        buf = io.BytesIO()
        document = SimpleDocTemplate(
            buf, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm,
            topMargin=17 * mm, bottomMargin=17 * mm,
            title=report.title or "模型推荐报告",
            author="银行模型市场智能推荐助手",
        )
        story: list[object] = [Paragraph(_inline_xml(report.title or "模型推荐报告"), title_style)]
        if report.summary:
            story.append(Paragraph(_inline_xml(report.summary), summary_style))
        story.append(Paragraph("<b>证据与使用边界：</b>" + html.escape(BOUNDARY_NOTICE), notice_style))

        for report_section in report.sections:
            story.append(Paragraph(_inline_xml(report_section.title), h1_style))
            pending_list: list[object] = []

            def flush_list() -> None:
                if pending_list:
                    story.append(ListFlowable(list(pending_list), bulletType="bullet", leftIndent=12, bulletFontName=font_name))
                    pending_list.clear()

            for block in _parse_blocks(report_section.content):
                if block.kind in {"bullet", "number"}:
                    pending_list.append(ListItem(Paragraph(_inline_xml(str(block.value)), body_style)))
                    continue
                flush_list()
                if block.kind == "heading":
                    story.append(Paragraph(_inline_xml(str(block.value)), h2_style))
                elif block.kind == "table":
                    rows = block.value
                    assert isinstance(rows, list)
                    column_count = max(len(row) for row in rows)
                    normalized = [row + [""] * (column_count - len(row)) for row in rows]
                    data = [[Paragraph(_inline_xml(cell), small_style) for cell in row] for row in normalized]
                    available = A4[0] - 36 * mm
                    if column_count == 5:
                        ratios = [0.07, 0.19, 0.14, 0.11, 0.49]
                    else:
                        ratios = [1 / column_count] * column_count
                    table = Table(data, colWidths=[available * ratio for ratio in ratios], repeatRows=1, hAlign="LEFT")
                    table.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCE6F1")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F2937")),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9CA3AF")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]))
                    story.extend([table, Spacer(1, 3 * mm)])
                else:
                    story.append(Paragraph(_inline_xml(str(block.value)), body_style))
            flush_list()

        story.append(Paragraph("报告元数据", h1_style))
        metadata = [
            ["报告编号", report.report_id],
            ["生成时间", report.generated_at],
            ["生成来源", report.generation_source],
        ]
        metadata_table = Table(
            [[Paragraph(html.escape(str(cell)), small_style) for cell in row] for row in metadata],
            colWidths=[32 * mm, A4[0] - 36 * mm - 32 * mm], hAlign="LEFT",
        )
        metadata_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F3F4F6")),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D1D5DB")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(metadata_table)

        def draw_footer(canvas: object, doc: object) -> None:
            canvas.saveState()
            canvas.setFont(font_name, 8)
            canvas.setFillColor(colors.HexColor("#6B7280"))
            canvas.drawCentredString(A4[0] / 2, 9 * mm, f"仅供人工复核 | 第 {doc.page} 页")
            canvas.restoreState()

        document.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)
        return buf.getvalue()


_report_exporter = ReportExporter()


def get_report_exporter() -> ReportExporter:
    return _report_exporter
